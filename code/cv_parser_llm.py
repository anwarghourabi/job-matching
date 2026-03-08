"""
CV Parser LLM — Version Groq (Gratuit & Ultra-Rapide) v4
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Gratuit, ~1 seconde par CV, aucune installation locale.

Prérequis :
  1. Créer un compte GRATUIT sur : https://console.groq.com
  2. Aller dans "API Keys" → "Create API Key"
  3. Copier la clé : gsk_XXXXXXXXXXXXXXXX

Configuration (une seule fois) :
  Option A — Variable d'environnement (recommandée) :
    Windows  : set GROQ_API_KEY=gsk_XXXXXXXX
    Mac/Linux: export GROQ_API_KEY=gsk_XXXXXXXX

  Option B — Dans le code :
    parser = CVParserLLM(api_key="gsk_XXXXXXXX")

Usage :
  parser = CVParserLLM()
  cv = parser.parse_file("mon_cv.pdf")
  cv = parser.parse_text("texte brut du CV...")
  print(cv)

Modèles Groq disponibles (gratuits) :
  - "llama-3.3-70b-versatile"   ← recommandé (très précis)
  - "llama-3.1-8b-instant"      ← plus rapide, moins précis
  - "mixtral-8x7b-32768"        ← bon pour les CV longs
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

import os
import re
import json
import subprocess
import sys
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional

# ── Installation automatique des dépendances ─────────────────
def _install(pkg):
    subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q"])

try:
    import pdfplumber
except ImportError:
    print("📦 Installation pdfplumber...")
    _install("pdfplumber")
    import pdfplumber

try:
    import docx
except ImportError:
    print("📦 Installation python-docx...")
    _install("python-docx")
    import docx

try:
    import requests
except ImportError:
    print("📦 Installation requests...")
    _install("requests")
    import requests


# ══════════════════════════════════════════════════════════════
# STRUCTURE DE SORTIE
# (identique à cv_parser.py pour compatibilité totale)
# ══════════════════════════════════════════════════════════════

@dataclass
class ParsedCV:
    raw_text:         str       = ""
    name:             str       = "Candidat"
    email:            str       = ""
    phone:            str       = ""
    summary:          str       = ""
    skills:           List[str] = field(default_factory=list)
    experience_level: str       = "unknown"  # entry / mid / senior / executive
    location:         str       = ""
    years_experience: int       = 0
    domain:           str       = "unknown"  # tech/finance/marketing/rh/sante/droit/management/vente/autre

    def is_valid(self) -> bool:
        return len(self.raw_text.strip()) > 50

    def __str__(self):
        return (
            f"👤 Nom          : {self.name}\n"
            f"📧 Email        : {self.email or 'N/A'}\n"
            f"📞 Téléphone    : {self.phone or 'N/A'}\n"
            f"📍 Localisation : {self.location or 'N/A'}\n"
            f"🎯 Domaine      : {self.domain}\n"
            f"👤 Niveau       : {self.experience_level}\n"
            f"📅 Années exp.  : {self.years_experience}\n"
            f"🔧 Skills ({len(self.skills)}): {', '.join(self.skills[:12])}\n"
            f"📝 Summary      : {self.summary[:200]}...\n"
        )


# ══════════════════════════════════════════════════════════════
# EXTRACTION TEXTE BRUT (PDF / DOCX / TXT)
# ══════════════════════════════════════════════════════════════

def _normalize_text(text: str) -> str:
    text = "".join(c if c.isprintable() or c in "\n\t" else " " for c in text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_pdf(path: str) -> str:
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text(x_tolerance=3, y_tolerance=3)
            if t:
                parts.append(t)
            else:
                words = page.extract_words()
                if words:
                    parts.append(" ".join(w["text"] for w in words))
    result = "\n".join(parts)
    if len(result.strip()) < 50:
        raise ValueError("PDF semble être une image scannée (aucun texte extractible).")
    return _normalize_text(result)


def extract_text_from_docx(path: str) -> str:
    d = docx.Document(path)
    lines = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return _normalize_text("\n".join(lines))


def extract_text_from_txt(path: str) -> str:
    for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            return _normalize_text(Path(path).read_text(encoding=enc))
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError("Impossible de décoder le fichier texte.")


def extract_raw_text(file_path: str) -> str:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Fichier introuvable : {file_path}")
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(str(path))
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(str(path))
    elif ext == ".txt":
        return extract_text_from_txt(str(path))
    else:
        raise ValueError(f"Format non supporté : {ext}. Utilise PDF, DOCX ou TXT.")


# ══════════════════════════════════════════════════════════════
# PROMPT
# ══════════════════════════════════════════════════════════════

SYSTEM_PROMPT = """You are an expert CV/resume parser. 
Extract structured information from any CV in any language (French, English, Arabic, or mixed).
Respond ONLY with a valid JSON object. No explanation, no markdown, no backticks. Just raw JSON.

JSON fields to extract:
- name: full name (string, "Candidat" if not found)
- email: email address (string, "" if absent)
- phone: phone number (string, "" if absent)
- location: city and country — be specific, e.g. "Paris, France" not just "Paris" (string, "" if absent)
- summary: 2-3 sentence professional summary describing the profile (string)
- skills: list of technical and professional skills relevant to the person's domain (array of strings, max 25)
- experience_level: MUST be exactly one of: "entry" / "mid" / "senior" / "executive"
  Rules — follow strictly:
  - "entry"     = 0-2 years of experience OR currently a student/intern seeking first job
  - "mid"       = 3-6 years of real professional experience
  - "senior"    = 7-14 years OR job title explicitly contains: Senior, Lead, Principal, Architect, Expert
  - "executive" = 15+ years OR C-level/top management title: CTO, CEO, CFO, COO, VP, Director, Head of
  ⚠️  CRITICAL RULE: A "Formation" section or a Master/Engineering degree does NOT mean the person
  is a student. Always base your decision on years of experience and current job title, not on education.
  Example: "Lead AI Engineer, 12 years experience, CentraleSupélec 2012" → "senior" NOT "entry"
- years_experience: total years of real professional experience as integer (0 if currently a student)
- domain: main professional domain — MUST be exactly one of:
  "tech" / "finance" / "marketing" / "rh" / "sante" / "droit" / "management" / "vente" / "autre"
"""

USER_PROMPT_TEMPLATE = """Parse this CV and return only the JSON object:

---
{cv_text}
---"""


# ══════════════════════════════════════════════════════════════
# CLIENT GROQ API
# ══════════════════════════════════════════════════════════════

GROQ_API_URL   = "https://api.groq.com/openai/v1/chat/completions"
DEFAULT_MODEL  = "llama-3.3-70b-versatile"


def _call_groq(cv_text: str, api_key: str, model: str) -> dict:
    """
    Appelle l'API Groq pour parser le CV.
    Retourne le dict JSON parsé.
    """
    # Groq supporte jusqu'à 32k tokens — on tronque à 12000 chars pour être safe
    cv_text_trimmed = cv_text[:12000] if len(cv_text) > 12000 else cv_text

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type":  "application/json",
    }

    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": USER_PROMPT_TEMPLATE.format(cv_text=cv_text_trimmed)},
        ],
        "temperature":    0.0,    # déterministe — crucial pour le parsing
        "max_tokens":     1024,
        "response_format": {"type": "json_object"},  # force JSON natif
    }

    try:
        response = requests.post(
            GROQ_API_URL,
            headers=headers,
            json=payload,
            timeout=30,
        )
    except requests.exceptions.ConnectionError:
        raise RuntimeError("❌ Pas de connexion internet. Vérifie ta connexion.")

    # Gestion des erreurs API
    if response.status_code == 401:
        raise RuntimeError(
            "❌ Clé API Groq invalide.\n"
            "   Vérifie ta clé sur : https://console.groq.com/keys"
        )
    if response.status_code == 429:
        raise RuntimeError(
            "❌ Limite de requêtes Groq atteinte.\n"
            "   Attends quelques secondes et réessaie.\n"
            "   Limite gratuite : 14 400 req/jour, 30 req/minute."
        )
    if response.status_code != 200:
        raise RuntimeError(
            f"❌ Erreur Groq ({response.status_code}) : {response.text[:200]}"
        )

    raw = response.json()["choices"][0]["message"]["content"].strip()

    # Nettoyer les éventuels artefacts markdown
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    raw = raw.strip()

    # Extraire le premier objet JSON valide si du texte parasite est présent
    json_match = re.search(r"\{.*\}", raw, re.DOTALL)
    if json_match:
        raw = json_match.group(0)

    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        raise ValueError(f"Réponse non-JSON de Groq : {raw[:300]}\nErreur : {e}")


# ══════════════════════════════════════════════════════════════
# PARSER PRINCIPAL
# ══════════════════════════════════════════════════════════════

class CVParserLLM:
    """
    Parser CV universel basé sur Groq API (gratuit).

    Avantages :
      ✅ Gratuit (14 400 requêtes/jour)
      ✅ Ultra-rapide (~1 seconde par CV)
      ✅ Aucune installation locale
      ✅ Tous profils, toutes langues, tous domaines
      ✅ Compatible avec matching_engine.py et api.py sans aucun changement

    Args:
        api_key : Clé API Groq (ou variable d'env GROQ_API_KEY)
        model   : Modèle Groq à utiliser (défaut: llama-3.3-70b-versatile)
        fallback: Si True, utilise cv_parser.py si Groq échoue
    """

    def __init__(
        self,
        api_key:  Optional[str] = None,
        model:    str           = DEFAULT_MODEL,
        fallback: bool          = True,
    ):
        self.api_key  = api_key or os.environ.get("GROQ_API_KEY", "")
        self.model    = model
        self.fallback = fallback

        if not self.api_key:
            raise ValueError(
                "\n❌ Clé API Groq manquante.\n\n"
                "   1. Crée un compte GRATUIT sur : https://console.groq.com\n"
                "   2. Va dans 'API Keys' → 'Create API Key'\n"
                "   3. Configure ta clé :\n\n"
                "   Windows  : set GROQ_API_KEY=gsk_XXXXXXXX\n"
                "   Mac/Linux: export GROQ_API_KEY=gsk_XXXXXXXX\n\n"
                "   Ou passe-la directement : CVParserLLM(api_key='gsk_XXXXXXXX')\n"
            )

        print(f"   ✅ CVParserLLM Groq prêt — modèle : {self.model}")

    def parse_text(self, raw_text: str) -> ParsedCV:
        """Parse un texte brut de CV"""
        raw_text = _normalize_text(raw_text)
        try:
            parsed = _call_groq(raw_text, self.api_key, self.model)
            return self._dict_to_parsed_cv(parsed, raw_text)
        except Exception as e:
            print(f"   ⚠️  Erreur Groq : {e}")
            if self.fallback:
                print("   🔄 Fallback vers cv_parser.py (règles)...")
                return self._fallback_parse(raw_text)
            raise

    def parse_file(self, file_path: str) -> ParsedCV:
        """Parse un fichier CV (PDF / DOCX / TXT)"""
        print(f"   📄 Lecture : {Path(file_path).name}")
        raw_text = extract_raw_text(file_path)
        if len(raw_text.strip()) < 50:
            raise ValueError("CV trop court ou illisible.")
        print(f"   ✅ Texte extrait : {len(raw_text):,} caractères")
        print(f"   🤖 Analyse Groq ({self.model}) en cours...")
        return self.parse_text(raw_text)

    def parse_text_manual(self, text: str, name: str = "Candidat") -> ParsedCV:
        """Parse avec un nom fourni manuellement"""
        cv = self.parse_text(text)
        if name and name != "Candidat":
            cv.name = name
        return cv

    # ── Validation et garde-fous ──────────────────────────────

    def _dict_to_parsed_cv(self, d: dict, raw_text: str) -> ParsedCV:
        """Convertit le dict JSON de Groq en ParsedCV avec validation"""

        valid_levels  = {"entry", "mid", "senior", "executive"}
        valid_domains = {
            "tech", "finance", "marketing", "rh",
            "sante", "droit", "management", "vente", "autre"
        }

        # Niveau
        exp_level = str(d.get("experience_level", "unknown")).lower().strip()
        if exp_level not in valid_levels:
            exp_level = "unknown"

        # Domaine
        domain = str(d.get("domain", "autre")).lower().strip()
        if domain not in valid_domains:
            domain = "autre"

        # Années
        years = d.get("years_experience", 0)
        try:
            years = int(years)
        except (ValueError, TypeError):
            years = 0

        # Skills
        skills = d.get("skills", [])
        if not isinstance(skills, list):
            skills = []
        skills = [str(s).strip() for s in skills if s and len(str(s).strip()) > 1]

        # ── Garde-fous : cohérence années ↔ niveau ────────────
        # Corrige les erreurs du LLM automatiquement
        if years >= 15:
            exp_level = "executive"
        elif years >= 10 and exp_level in ("entry", "mid"):
            exp_level = "senior"
        elif years >= 7 and exp_level == "entry":
            exp_level = "senior"
        elif years >= 3 and exp_level == "entry":
            # Vérifier si c'est vraiment un étudiant
            student_signals = [
                "stage pfe", "pfe", "looking for internship",
                "currently studying", "en quête d'un stage",
                "3ème année", "2ème année", "1ère année",
                "cherche stage", "recherche stage",
                "final year student", "undergraduate",
            ]
            is_student = any(s in raw_text.lower() for s in student_signals)
            if not is_student:
                exp_level = "mid"

        return ParsedCV(
            raw_text         = raw_text,
            name             = str(d.get("name",    "Candidat")).strip() or "Candidat",
            email            = str(d.get("email",   "")).strip(),
            phone            = str(d.get("phone",   "")).strip(),
            location         = str(d.get("location","")).strip(),
            summary          = str(d.get("summary", "")).strip(),
            skills           = skills,
            experience_level = exp_level,
            years_experience = max(0, years),
            domain           = domain,
        )

    def _fallback_parse(self, raw_text: str) -> ParsedCV:
        """Fallback vers cv_parser.py si Groq est indisponible"""
        try:
            from cv_parser import CVParser as _CVParser
            print("   ✅ Fallback cv_parser.py OK")
            return _CVParser().parse_text(raw_text)
        except ImportError:
            print("   ⚠️  cv_parser.py non trouvé — retour minimal")
            return ParsedCV(raw_text=raw_text, summary=raw_text[:400])


# ══════════════════════════════════════════════════════════════
# ALIAS — compatibilité totale avec matching_engine.py et api.py
# ══════════════════════════════════════════════════════════════
CVParser = CVParserLLM


# ══════════════════════════════════════════════════════════════
# TEST STANDALONE
# ══════════════════════════════════════════════════════════════

if __name__ == "__main__":

    API_KEY = os.environ.get("GROQ_API_KEY", "")
    if not API_KEY:
        print("\n❌ Configure ta clé Groq d'abord :")
        print("   Windows  : set GROQ_API_KEY=gsk_XXXXXXXX")
        print("   Mac/Linux: export GROQ_API_KEY=gsk_XXXXXXXX")
        sys.exit(1)

    print("\n" + "="*65)
    print("TEST CV PARSER — Groq API v4")
    print("="*65)

    parser = CVParserLLM(api_key=API_KEY)

    tests = [
        (
            "Alexandre Moreau — Senior AI (le cas problématique)",
            """Alexandre Moreau — Senior AI / Machine Learning Engineer
            Paris, France | alexandre.moreau@email.com | +33 6 12 34 56 78
            12 ans d'expérience. Lead AI Engineer chez Dataflow Technologies.
            Expert LLMs, Computer Vision, MLOps. CentraleSupélec 2012.
            Formation : Diplôme ingénieur CentraleSupélec 2010-2012.
            Compétences : Python, PyTorch, TensorFlow, AWS SageMaker, Kubernetes, MLflow.""",
            "senior",
        ),
        (
            "Ghada — Étudiante BI Tunisie",
            """3ème année licence Business Intelligence à l'ISG-Tunis.
            Je suis en quête d'un stage PFE.
            Compétences : Power BI, Talend, SQL, Python, Machine Learning, MDX.""",
            "entry",
        ),
        (
            "Pierre — Comptable confirmé",
            """Pierre Dupont — Comptable confirmé, 5 ans d'expérience. Paris, France.
            Comptabilité générale, Audit, IFRS, Sage, Excel, VBA, Budget, Trésorerie.""",
            "mid",
        ),
        (
            "John Smith — HR Director",
            """HR Director with 15 years of experience. London, UK.
            Led HR teams of 20+ across 5 countries.
            Recruitment, Payroll, Workday HRIS, Employment Law, D&I.""",
            "executive",
        ),
    ]

    passed = 0
    for label, cv_text, expected in tests:
        print(f"\n📋 {label}")
        print("─" * 60)
        cv = parser.parse_text(cv_text)
        print(cv)
        ok = cv.experience_level == expected
        print(f"{'✅' if ok else '❌'} Niveau = '{cv.experience_level}' (attendu: '{expected}')")
        if ok:
            passed += 1

    print(f"\n{'='*65}")
    print(f"Résultats : {passed}/{len(tests)} tests réussis")
    if passed == len(tests):
        print("✅ CVParserLLM Groq — 100% opérationnel !")
    else:
        print("⚠️  Vérifier les réponses du modèle.")
    print("="*65)